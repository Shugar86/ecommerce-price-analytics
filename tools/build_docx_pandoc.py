#!/usr/bin/env python3
"""Build the VKR DOCX with Pandoc and a GOST-oriented reference template.

The script keeps the source document in Markdown, converts legacy bold
structural headings to real Markdown headings in a temporary file, and lets
Pandoc convert LaTeX math blocks to native editable Word equations.
Table content is produced entirely by Pandoc; postprocessing only adjusts
visual styling without rewriting table grid XML.
"""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, Twips
from docx.table import _Cell


ROOT = Path(__file__).resolve().parent.parent
DEFAULT_INPUT = ROOT / "VKR_FINAL.md"
DEFAULT_OUTPUT = ROOT / "VKR_FINAL.docx"
DEFAULT_REFERENCE = ROOT / "assets" / "docx" / "reference_gost.docx"
DEFAULT_OUTPUT_COPY = ROOT / "output" / "VKR_FINAL_готовая_версия_pandoc.docx"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_NS = {"w": W_NS}


def ensure_pandoc() -> None:
    """Raise a clear error if Pandoc is unavailable."""
    if shutil.which("pandoc") is None:
        raise RuntimeError("Pandoc не найден. Установите pandoc 3.x и повторите сборку.")


def configure_style(
    style,
    *,
    size_pt: int = 14,
    bold: bool = False,
    align=None,
    keep_with_next: bool = False,
    keep_together: bool = False,
    all_caps: bool = False,
) -> None:
    """Apply Times New Roman and basic paragraph settings to a Word style."""
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    if all_caps:
        style.font.all_caps = True
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.keep_with_next = keep_with_next
    pf.keep_together = keep_together
    if align is not None:
        pf.alignment = align


def add_page_number_footer(doc) -> None:
    """Insert a centred PAGE field in the footer with title-page suppression."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    sect_pr = section._sectPr
    title_pg = sect_pr.find(qn("w:titlePg"))
    if title_pg is None:
        title_pg = OxmlElement("w:titlePg")
        sect_pr.append(title_pg)

    footer = section.footer
    if footer.paragraphs:
        para = footer.paragraphs[0]
    else:
        para = footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_instr = OxmlElement("w:instrText")
    fld_instr.text = "PAGE   \\* MERGEFORMAT"
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin)
    run._element.append(fld_instr)
    run._element.append(fld_separate)
    run._element.append(fld_end)


def create_reference_docx(path: Path) -> None:
    """Create a DOCX reference template with GOST-like page and text styles."""
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)

    normal = doc.styles["Normal"]
    configure_style(normal, size_pt=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    normal.paragraph_format.first_line_indent = Cm(1.25)

    for style_name in ("Body Text", "First Paragraph"):
        if style_name in doc.styles:
            style = doc.styles[style_name]
        else:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = normal
        configure_style(style, size_pt=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        style.paragraph_format.first_line_indent = Cm(1.25)

    for style_name, size, alignment, caps in (
        ("Title", 14, WD_ALIGN_PARAGRAPH.CENTER, True),
        ("Heading 1", 14, WD_ALIGN_PARAGRAPH.CENTER, True),
        ("Heading 2", 14, WD_ALIGN_PARAGRAPH.LEFT, False),
        ("Heading 3", 14, WD_ALIGN_PARAGRAPH.LEFT, False),
        ("Heading 4", 14, WD_ALIGN_PARAGRAPH.LEFT, False),
    ):
        style = doc.styles[style_name]
        configure_style(
            style,
            size_pt=size,
            bold=True,
            align=alignment,
            keep_with_next=True,
            keep_together=True,
            all_caps=caps,
        )
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(6 if style_name != "Heading 1" else 12)
        style.paragraph_format.space_after = Pt(6 if style_name != "Heading 1" else 12)

    for style_name in ("Caption", "Image Caption"):
        if style_name in doc.styles:
            style = doc.styles[style_name]
        else:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = normal
        configure_style(
            style,
            size_pt=14,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            keep_with_next=True,
            keep_together=True,
        )
        style.paragraph_format.first_line_indent = Cm(0)

    for style_name in ("List Bullet", "List Number"):
        if style_name in doc.styles:
            style = doc.styles[style_name]
            configure_style(style, size_pt=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            style.paragraph_format.first_line_indent = Cm(0)

    add_page_number_footer(doc)

    doc.add_paragraph("GOST reference template for Pandoc.").style = normal
    doc.save(path)


def _cell_shading(cell, fill: str) -> None:
    """Заливает ячейку таблицы цветом ``fill`` (hex без #)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _row_is_header(row) -> bool:
    """True, если строка помечена Pandoc/Word как заголовок таблицы."""
    tr_pr = row._tr.find(qn("w:trPr"))
    if tr_pr is None:
        return False
    return tr_pr.find(qn("w:tblHeader")) is not None


def _row_effective_cols(tr) -> int:
    """Считает ширину строки таблицы в колонках сетки с учётом gridSpan."""
    total = 0
    for tc in tr.findall(f"{{{W_NS}}}tc"):
        span = 1
        tc_pr = tc.find(f"{{{W_NS}}}tcPr")
        if tc_pr is not None:
            grid_span = tc_pr.find(f"{{{W_NS}}}gridSpan")
            if grid_span is not None:
                span = int(grid_span.get(f"{{{W_NS}}}val", "1"))
        total += span
    return total


def _set_table_full_width(table, content_width: Twips) -> None:
    """Растягивает таблицу на 100 % ширины без перезаписи tblGrid."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "autofit")

    try:
        table.width = content_width
    except (AttributeError, ValueError):
        pass


def _set_cell_wrap(cell) -> None:
    """Разрешает перенос текста внутри ячейки."""
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap = tc_pr.find(qn("w:noWrap"))
    if no_wrap is not None:
        tc_pr.remove(no_wrap)

    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        side_el = tc_mar.find(qn(f"w:{side}"))
        if side_el is None:
            side_el = OxmlElement(f"w:{side}")
            tc_mar.append(side_el)
        side_el.set(qn("w:w"), "57")
        side_el.set(qn("w:type"), "dxa")


def _ensure_run_font(run, *, font_pt: int, header: bool) -> None:
    """Задаёт явный шрифт run, чтобы Word не терял текст из-за стиля Compact."""
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(font_pt)
    run.font.color.rgb = None
    if header:
        run.bold = True


def _format_table_cell(cell, *, header: bool, font_pt: int) -> None:
    """Приводит ячейку к компактному табличному виду по ГОСТ."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _set_cell_wrap(cell)
    for paragraph in cell.paragraphs:
        try:
            paragraph.style = "Normal"
        except KeyError:
            pass
        p_pr = paragraph._element.get_or_add_pPr()
        p_style = p_pr.find(qn("w:pStyle"))
        if p_style is not None:
            p_pr.remove(p_style)
        pf = paragraph.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(0)
        pf.right_indent = Cm(0)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.0
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if paragraph.runs:
            for run in paragraph.runs:
                _ensure_run_font(run, font_pt=font_pt, header=header)
        else:
            plain = paragraph.text
            if plain:
                paragraph.clear()
                run = paragraph.add_run(plain)
                _ensure_run_font(run, font_pt=font_pt, header=header)
    if header:
        _cell_shading(cell, "D9D9D9")


def _format_table_captions(doc: Document) -> None:
    """Оформляет абзацы «Таблица N — …» перед таблицами."""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text.startswith("Таблица "):
            continue
        pf = paragraph.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.keep_with_next = True
        pf.space_before = Pt(6)
        pf.space_after = Pt(3)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(12)
            run.bold = True


def _strip_table_cells(row: str) -> list[str]:
    """Разбирает строку pipe-таблицы на ячейки."""
    row = row.strip()
    if row.startswith("|"):
        row = row[1:]
    if row.endswith("|"):
        row = row[:-1]
    return [cell.strip() for cell in row.split("|")]


def _is_table_separator(parts: list[str]) -> bool:
    """True для строки-разделителя GFM ``| :--- | ---: |``."""
    cells = [part.strip() for part in parts if part.strip()]
    if not cells:
        return False
    for cell in cells:
        if not re.match(r"^:?-{2,}:?$", cell.replace(" ", "")):
            return False
    return True


def _normalize_table_cell_text(text: str) -> str:
    """Убирает markdown/latex-обёртки в ячейках и подписях таблиц."""
    if not text:
        return text
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\$([^$]+)\$", r"\1", text)
    text = re.sub(r"\\text\{([^}]+)\}", r"\1", text)
    text = text.replace("\\_", "_")
    return text.strip()


def validate_docx_tables(docx_path: Path) -> list[str]:
    """Проверяет целостность сетки таблиц в готовом DOCX."""
    issues: list[str] = []
    root = ET.fromstring(zipfile.ZipFile(docx_path).read("word/document.xml"))
    for ti, tbl in enumerate(root.findall(".//w:tbl", XML_NS)):
        tbl_grid = tbl.find("w:tblGrid", XML_NS)
        if tbl_grid is None:
            issues.append(f"Таблица {ti + 1}: отсутствует tblGrid")
            continue
        grid_cols = len(tbl_grid.findall("w:gridCol", XML_NS))
        if grid_cols == 0:
            issues.append(f"Таблица {ti + 1}: пустой tblGrid")
            continue
        for ri, tr in enumerate(tbl.findall("w:tr", XML_NS)):
            effective = _row_effective_cols(tr)
            if effective != grid_cols:
                cap = ""
                first_tc = tr.find("w:tc", XML_NS)
                if first_tc is not None:
                    cap = "".join(t.text or "" for t in first_tc.findall(".//w:t", XML_NS))[:40]
                issues.append(
                    f"Таблица {ti + 1}, строка {ri + 1}: сетка {grid_cols}, фактически {effective} ({cap!r})"
                )
    return issues


def postprocess_docx_tables(docx_path: Path) -> None:
    """Исправляет таблицы после Pandoc: сетка, ширина, шрифт, шапка."""
    doc = Document(str(docx_path))
    section = doc.sections[0]
    content_width = section.page_width - section.left_margin - section.right_margin

    _format_table_captions(doc)

    for table in doc.tables:
        ncols = len(table.columns)
        font_pt = 10 if ncols >= 7 else 11 if ncols >= 5 else 12
        try:
            table.style = "Table Grid"
        except KeyError:
            table.style = "Table Normal"

        _set_table_full_width(table, content_width)

        for ri, row in enumerate(table.rows):
            header = _row_is_header(row) or ri == 0
            # row.cells у python-docx может возвращать «чужие» tc после Pandoc;
            # обходим XML-узлы строки напрямую.
            for tc in row._tr.findall(qn("w:tc")):
                _format_table_cell(_Cell(tc, table), header=header, font_pt=font_pt)

    doc.save(str(docx_path))


def structural_heading_level(title: str) -> int | None:
    """Return Markdown heading level for bold structural headings."""
    if (
        title == "ВВЕДЕНИЕ"
        or title.startswith("ГЛАВА ")
        or title in {"ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", "ПРИЛОЖЕНИЯ"}
        or title.startswith("Приложение ")
    ):
        return 1
    if title and title[0].isdigit():
        leading = title.split(" ", maxsplit=1)[0]
        numeric_parts = leading.split(".")
        if all(part.isdigit() for part in numeric_parts):
            depth = len(numeric_parts)
            if depth == 1:
                return 1
            if depth == 2:
                return 2
            if depth >= 3:
                return 3
    return None


def _normalize_display_math_tags(text: str) -> str:
    """Заменяет ``\\tag{n.m}`` на ``\\qquad (n.m)`` для корректного OMML в Word."""
    return re.sub(r"\\tag\{([^}]+)\}", r"\\qquad (\1)", text)


def preprocess_markdown(source: Path, target: Path) -> None:
    """Convert local Markdown conventions to Pandoc-friendly Markdown."""
    output_lines: list[str] = []
    in_math = False
    for line in source.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if stripped == "$$":
            in_math = not in_math
            output_lines.append(line)
            continue
        if in_math:
            output_lines.append(_normalize_display_math_tags(line))
            continue
        if stripped.startswith("Таблица "):
            output_lines.append(_normalize_table_cell_text(stripped))
            continue
        if stripped.startswith("|") and "|" in stripped:
            cells = _strip_table_cells(line)
            if not _is_table_separator(cells):
                normalized_cells = []
                for ci, cell in enumerate(cells):
                    value = _normalize_table_cell_text(cell)
                    if not value and ci == 0:
                        value = " "
                    normalized_cells.append(value)
                output_lines.append("| " + " | ".join(normalized_cells) + " |")
            else:
                output_lines.append(line)
            continue
        if stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            title = stripped[2:-2].strip()
            level = structural_heading_level(title)
            if level is not None:
                output_lines.append(f"{'#' * level} {title}")
                continue
        output_lines.append(line)
    target.write_text("\n".join(output_lines) + "\n", encoding="utf-8")


def run_pandoc(input_md: Path, output_docx: Path, reference_docx: Path) -> None:
    """Run Pandoc with settings suitable for a thesis-style DOCX."""
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    cmd = [
        "pandoc",
        str(input_md),
        "--from=markdown+tex_math_dollars+pipe_tables+fenced_code_blocks",
        "--to=docx",
        f"--reference-doc={reference_docx}",
        "--resource-path=.",
        "--wrap=none",
        "--output",
        str(output_docx),
    ]
    subprocess.run(cmd, cwd=ROOT, check=True)


def build_docx(source: Path, output: Path, reference: Path, output_copy: Path | None) -> None:
    """Build the main DOCX and optional output copy."""
    ensure_pandoc()
    if not source.is_file():
        raise FileNotFoundError(f"Markdown-файл не найден: {source}")
    if reference.is_file():
        reference.unlink()
    create_reference_docx(reference)

    with tempfile.TemporaryDirectory(prefix="vkr-pandoc-") as temp_dir:
        prepared = Path(temp_dir) / "VKR_FINAL.pandoc.md"
        preprocess_markdown(source, prepared)
        run_pandoc(prepared, output, reference)

    postprocess_docx_tables(output)

    table_issues = validate_docx_tables(output)
    if table_issues:
        details = "\n".join(f"  - {issue}" for issue in table_issues[:10])
        raise RuntimeError(
            f"DOCX содержит {len(table_issues)} ошибок структуры таблиц:\n{details}"
        )

    if output_copy is not None:
        output_copy.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(output, output_copy)


def main() -> int:
    """CLI entrypoint."""
    parser = argparse.ArgumentParser(description="Build VKR DOCX with Pandoc and GOST styles.")
    parser.add_argument("--input", type=Path, default=DEFAULT_INPUT)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--reference-doc", type=Path, default=DEFAULT_REFERENCE)
    parser.add_argument("--copy-to-output", type=Path, default=DEFAULT_OUTPUT_COPY)
    args = parser.parse_args()

    try:
        build_docx(args.input, args.output, args.reference_doc, args.copy_to_output)
    except (FileNotFoundError, RuntimeError, subprocess.CalledProcessError) as exc:
        print(f"Ошибка сборки DOCX через Pandoc: {exc}", file=sys.stderr)
        return 1

    print(f"Создан: {args.output}")
    if args.copy_to_output:
        print(f"Копия: {args.copy_to_output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
