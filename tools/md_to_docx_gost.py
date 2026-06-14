#!/usr/bin/env python3
"""
Конвертация Markdown ВКР в DOCX с базовым оформлением по ГОСТ 7.32 (отчёт о НИР / ВКР).

Ориентиры (типичные для вузов): Times New Roman 14 пт, интервал 1,5; абзац 1,25 см;
поля: левое 30 мм (переплёт), правое 15 мм, верх/низ 20 мм; выравнивание по ширине.

Запуск: python tools/md_to_docx_gost.py [--input path] [--output path]
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Mm, Pt


def _set_run_font(run, name: str = "Times New Roman", size_pt: int = 14, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def _apply_gost_page(doc: Document) -> None:
    sec = doc.sections[0]
    sec.left_margin = Mm(30)
    sec.right_margin = Mm(15)
    sec.top_margin = Mm(20)
    sec.bottom_margin = Mm(20)


def _setup_styles(doc: Document) -> None:
    """Настройка Normal и пользовательских стилей заголовков."""
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    nf = normal.paragraph_format
    nf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    nf.line_spacing = 1.5
    nf.first_line_indent = Cm(1.25)
    nf.space_after = Pt(0)
    nf.space_before = Pt(0)
    nf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Heading 1 — название главы
    if "GOST Heading 1" not in [s.name for s in doc.styles]:
        h1 = doc.styles.add_style("GOST Heading 1", WD_STYLE_TYPE.PARAGRAPH)
    else:
        h1 = doc.styles["GOST Heading 1"]
    h1.base_style = doc.styles["Normal"]
    h1.font.bold = True
    h1.font.size = Pt(16)
    h1.font.name = "Times New Roman"
    h1.paragraph_format.first_line_indent = Cm(0)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for name, size, indent_cm in (
        ("GOST Heading 2", Pt(14), 0),
        ("GOST Heading 3", Pt(14), 0),
    ):
        if name not in [s.name for s in doc.styles]:
            st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            st = doc.styles[name]
        st.base_style = doc.styles["Normal"]
        st.font.bold = True
        st.font.size = size
        st.paragraph_format.first_line_indent = Cm(indent_cm)
        st.paragraph_format.space_before = Pt(6)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_inline(paragraph, text: str, font_pt: int = 14) -> None:
    """Добавляет runs с **жирным** и `моноширинным`."""
    if not text:
        return
    text = _normalize_typography(text)
    # Split by **bold** and `code`
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _set_run_font(run, size_pt=font_pt, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
            run.font.size = Pt(font_pt - 1)
        else:
            run = paragraph.add_run(part)
            _set_run_font(run, size_pt=font_pt)


def _normalize_typography(text: str) -> str:
    """Нормализует тире и диапазоны годов для аккуратной типографики в DOCX."""
    if not text:
        return text
    normalized = re.sub(r"(?<=\S)\s+---\s+(?=\S)", " — ", text)
    normalized = re.sub(r"(?<=\D)(\d{4})--(\d{4})(?=\D|$)", r"\1–\2", normalized)
    return normalized


def _add_paragraph(doc: Document, text: str, no_indent: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if no_indent:
        p.paragraph_format.first_line_indent = Cm(0)
    _add_inline(p, text)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5


def _structural_heading_level(text: str) -> int | None:
    """Возвращает уровень заголовка для строк вида **ВВЕДЕНИЕ** или **1.2. Раздел**."""
    if not (text.startswith("**") and text.endswith("**")):
        return None
    title = text[2:-2].strip()
    if (
        title == "ВВЕДЕНИЕ"
        or title.startswith("ГЛАВА ")
        or title in {"ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", "ПРИЛОЖЕНИЯ"}
        or title.startswith("Приложение ")
    ):
        return 1
    if re.match(r"^\d+\.\d+\.\d+\.", title):
        return 3
    if re.match(r"^\d+\.\d+\.", title):
        return 2
    return None


def _add_structural_heading(doc: Document, raw: str, level: int) -> None:
    """Добавляет структурный заголовок с жирным шрифтом и без абзацного отступа."""
    title = raw[2:-2].strip()
    style = "GOST Heading 1" if level == 1 else "GOST Heading 2" if level == 2 else "GOST Heading 3"
    p = doc.add_paragraph(style=style)
    r = p.add_run(title)
    _set_run_font(r, size_pt=16 if level == 1 else 14, bold=True)


def _add_image(doc: Document, md_path: Path, alt: str, src: str) -> None:
    """Вставляет изображение из Markdown-ссылки, если файл доступен."""
    image_path = Path(src)
    if not image_path.is_absolute():
        image_path = (md_path.parent / image_path).resolve()
    if not image_path.is_file():
        _add_paragraph(doc, f"[Изображение не найдено: {src}]", no_indent=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(15.5))


def _add_code_block(doc: Document, lines: list[str], lang: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    label = f"[{lang}] " if lang and lang != "text" else ""
    run = p.add_run(label + "\n".join(lines))
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(10)


def _add_math_block(doc: Document, lines: list[str]) -> None:
    """Добавляет блочную формулу из Markdown $$...$$ без служебных разделителей."""
    formula = " ".join(line.strip() for line in lines if line.strip())
    if not formula:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(formula)
    run.font.name = "Cambria Math"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    run.font.size = Pt(14)


def _cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Table Grid"
    for ri, row_cells in enumerate(rows):
        for ci in range(ncols):
            cell = tbl.rows[ri].cells[ci]
            txt = row_cells[ci] if ci < len(row_cells) else ""
            p = cell.paragraphs[0]
            p.text = ""
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.0
            _add_inline(p, txt.strip(), font_pt=12)
            for run in p.runs:
                run.font.size = Pt(12)
            if ri == 0:
                _cell_shading(cell, "D9D9D9")


def _strip_table_cells(row: str) -> list[str]:
    row = row.strip()
    if row.startswith("|"):
        row = row[1:]
    if row.endswith("|"):
        row = row[:-1]
    return [c.strip() for c in row.split("|")]


def _is_table_separator(parts: list[str]) -> bool:
    if not parts:
        return False
    cells = [p.strip() for p in parts if p.strip()]
    if not cells:
        return False
    for c in cells:
        # GFM: :---, ---, :-----:, -:-
        if not re.match(r"^:?-{2,}:?$", c.replace(" ", "")):
            return False
    return True


def _is_simple_table_border(line: str) -> bool:
    """Проверяет строку-границу ASCII-таблицы вида '-----'."""
    return bool(re.match(r"^\s*-{10,}\s*$", line))


def _is_simple_table_separator(line: str) -> bool:
    """Проверяет строку-разделитель колонок в простой ASCII-таблице."""
    stripped = line.strip()
    return bool(re.match(r"^-{3,}(?:\s+-{3,})+\s*$", stripped))


def _split_simple_table_row(row_line: str, ncols: int) -> list[str]:
    """Разбивает строку ASCII-таблицы на ячейки по 2+ пробелам."""
    parts = [c.strip() for c in re.split(r"\s{2,}", row_line.strip(), maxsplit=max(0, ncols - 1))]
    if len(parts) < ncols:
        parts.extend([""] * (ncols - len(parts)))
    return parts[:ncols]


def _parse_simple_ascii_table(lines: list[str], start_idx: int) -> tuple[list[list[str]], int] | None:
    """Парсит таблицу в простом текстовом формате (Pandoc simple table)."""
    if not _is_simple_table_border(lines[start_idx]):
        return None

    i = start_idx + 1
    body: list[str] = []
    while i < len(lines) and not _is_simple_table_border(lines[i]):
        body.append(lines[i].rstrip())
        i += 1
    if i >= len(lines):
        return None

    non_empty = [ln for ln in body if ln.strip()]
    if len(non_empty) < 3:
        return None

    separator_pos: int | None = None
    for idx, row in enumerate(non_empty):
        if _is_simple_table_separator(row):
            separator_pos = idx
            break
    if separator_pos is None or separator_pos == 0:
        return None

    rows: list[list[str]] = []
    header_split = [c.strip() for c in re.split(r"\s{2,}", non_empty[0].strip()) if c.strip()]
    ncols = len(header_split)
    if ncols <= 1:
        ncols = len(re.findall(r"-{3,}", non_empty[separator_pos]))
    if ncols <= 1:
        return None

    header_cells = _split_simple_table_row(non_empty[0], ncols)
    rows.append(header_cells)

    for row in non_empty[separator_pos + 1 :]:
        if _is_simple_table_separator(row):
            continue
        cells = _split_simple_table_row(row, ncols)
        if any(cell for cell in cells):
            rows.append(cells)

    if len(rows) <= 1:
        return None
    return rows, i + 1


def convert_md_to_docx(md_path: Path, out_path: Path) -> None:
    raw = md_path.read_text(encoding="utf-8")
    lines = raw.splitlines()

    doc = Document()
    _apply_gost_page(doc)
    _setup_styles(doc)

    i = 0
    while i < len(lines):
        line = lines[i]

        # Пустая строка
        if not line.strip():
            i += 1
            continue

        # Структурные заголовки в текущем Markdown оформлены как **...**.
        heading_level = _structural_heading_level(line.strip())
        if heading_level is not None:
            _add_structural_heading(doc, line.strip(), heading_level)
            i += 1
            continue

        # Простая ASCII-таблица (строки с границей из дефисов и колонками).
        ascii_table = _parse_simple_ascii_table(lines, i)
        if ascii_table is not None:
            table_rows, next_idx = ascii_table
            _add_table(doc, table_rows)
            i = next_idx
            continue

        # Горизонтальная линия / разделитель
        if re.match(r"^[-*_]{3,}\s*$", line.strip()):
            i += 1
            continue

        # Изображения Markdown: ![подпись](path/to/image.png)
        m_img = re.match(r"^!\[(.*?)\]\((.*?)\)\s*$", line.strip())
        if m_img:
            _add_image(doc, md_path, m_img.group(1), m_img.group(2))
            i += 1
            continue

        # Блок кода
        if line.strip().startswith("```"):
            lang = line.strip()[3:].strip() or "text"
            chunk: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                chunk.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            _add_code_block(doc, chunk, lang)
            continue

        # Блочные формулы Markdown/LaTeX: $$ ... $$.
        if line.strip() == "$$":
            chunk: list[str] = []
            i += 1
            while i < len(lines) and lines[i].strip() != "$$":
                chunk.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            _add_math_block(doc, chunk)
            continue

        # Таблица
        if "|" in line and line.strip().startswith("|"):
            table_rows: list[list[str]] = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                parts = _strip_table_cells(lines[i])
                if _is_table_separator(parts):
                    i += 1
                    continue
                table_rows.append(parts)
                i += 1
            _add_table(doc, table_rows)
            continue

        # Заголовки
        if line.startswith("# "):
            t = line[2:].strip()
            p = doc.add_paragraph(style="GOST Heading 1")
            r = p.add_run(t)
            _set_run_font(r, size_pt=16, bold=True)
            i += 1
            continue
        if line.startswith("## "):
            t = line[3:].strip()
            p = doc.add_paragraph(style="GOST Heading 2")
            r = p.add_run(t)
            _set_run_font(r, size_pt=14, bold=True)
            i += 1
            continue
        if line.startswith("### "):
            t = line[4:].strip()
            p = doc.add_paragraph(style="GOST Heading 3")
            r = p.add_run(t)
            _set_run_font(r, size_pt=14, bold=True)
            i += 1
            continue

        # Маркированный / нумерованный список
        # Markdown from LLM often escapes ordered-list dots as ``1\.``.
        # Treat both ``1. text`` and ``1\. text`` as the same list item, so
        # DOCX/PDF export does not produce artifacts like "11.".
        m_num = re.match(r"^(\s*)(\d+)\\?[.)]\s+(.*)$", line)
        m_bul = re.match(r"^(\s*)[-*+]\s+(.*)$", line)
        m_bul_quote = re.match(r"^\s*>\s*[-]{1,3}\s+(.*)$", line)
        if m_num:
            indent_spaces = len(m_num.group(1).replace("\t", "    "))
            level = indent_spaces // 2
            txt = m_num.group(3)
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0.8 + level * 0.5)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(f"{m_num.group(2)}. ")
            _set_run_font(run, size_pt=14)
            _add_inline(p, txt)
            i += 1
            continue
        if m_bul:
            indent_spaces = len(m_bul.group(1).replace("\t", "    "))
            level = indent_spaces // 2
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
            _add_inline(p, m_bul.group(2))
            i += 1
            continue
        if m_bul_quote:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0.6)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
            _add_inline(p, m_bul_quote.group(1))
            i += 1
            continue

        _add_paragraph(doc, line.strip())
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> int:
    ap = argparse.ArgumentParser(description="MD → DOCX (оформление по ГОСТ 7.32, базовое)")
    ap.add_argument(
        "--input",
        type=Path,
        default=Path(__file__).resolve().parent.parent / "VKR_FINAL.md",
    )
    ap.add_argument(
        "--output",
        type=Path,
        default=Path(__file__).resolve().parent.parent / "VKR_FINAL.docx",
    )
    args = ap.parse_args()
    if not args.input.is_file():
        print(f"Файл не найден: {args.input}", file=sys.stderr)
        return 1
    convert_md_to_docx(args.input, args.output)
    print(f"Создан: {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
